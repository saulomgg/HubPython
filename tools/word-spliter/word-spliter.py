import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
import math

input_docx = None  # Variável global para armazenar o arquivo DOCX selecionado
output_directory = ""  # Variável global para armazenar a pasta de saída

def split(input_docx, output_directory, num_parts):
    doc = Document(input_docx)
    total_pages = len(doc.element.xpath('//w:sectPr'))
    file_size = os.path.getsize(input_docx) / (1024 * 1024)  # Tamanho em MB

    pages_per_part = math.ceil(total_pages / num_parts)

    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    for part in range(num_parts):
        start_page = part * pages_per_part
        end_page = min((part + 1) * pages_per_part, total_pages)

        part_doc = Document()
        part_doc.element.body.extend(doc.element.body[start_page:end_page])
        
        output_docx = os.path.join(output_directory, f"part_{part + 1}.docx")
        part_doc.save(output_docx)

        pages_in_part = end_page - start_page
        print(f"Parte {part + 1} criada: {output_docx}")
        print(f"Páginas na parte: {pages_in_part}")

    # Exibir uma mensagem informativa
    messagebox.showinfo("Concluído", f"Documento dividido com sucesso em {num_parts} partes.\nNúmero de páginas: {total_pages}\nTamanho do arquivo: {file_size:.2f} MB")

def select_file():
    global input_docx
    input_docx = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if input_docx:
        label_info.config(text=f"Arquivo selecionado: {input_docx}")
        doc = Document(input_docx)
        total_pages = len(doc.element.xpath('//w:sectPr'))
        file_size = os.path.getsize(input_docx) / (1024 * 1024)  # Tamanho em MB
        label_pages.config(text=f"Páginas no DOCX: {total_pages}")
        label_size.config(text=f"Tamanho do arquivo: {file_size:.2f} MB")
        update_parts_info()

def select_output_directory():
    global output_directory
    output_directory = filedialog.askdirectory(title="Selecione a pasta de destino para as partes")
    if output_directory:
        label_output_directory.config(text=f"Pasta de saída selecionada: {output_directory}")
        update_parts_info()

def update_parts_info():
    global input_docx, output_directory
    num_parts = int(entry_parts.get())

    if input_docx and num_parts > 0 and output_directory:
        doc = Document(input_docx)
        total_pages = len(doc.element.xpath('//w:sectPr'))
        pages_per_part = math.ceil(total_pages / num_parts)

        info_text = f"Páginas por parte: {pages_per_part}"
        label_parts_info.config(text=info_text)
    else:
        label_parts_info.config(text="")

# Configuração da janela principal
root = tk.Tk()
root.title("Word Splitter")

# Configurar a estrutura visual com uma grade
frame = ttk.Frame(root, padding=10)
frame.grid(column=0, row=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Botão para selecionar o arquivo DOCX
btn_select_file = ttk.Button(frame, text="Selecionar DOCX", command=select_file)
btn_select_file.grid(row=0, column=0, padx=5, pady=5)

# Rótulo para exibir informações sobre o arquivo DOCX
label_info = ttk.Label(frame, text="")
label_info.grid(row=1, column=0, padx=5, pady=5)

# Rótulo para exibir o número de páginas no DOCX
label_pages = ttk.Label(frame, text="")
label_pages.grid(row=2, column=0, padx=5, pady=5)

# Rótulo para exibir o tamanho do arquivo DOCX
label_size = ttk.Label(frame, text="")
label_size.grid(row=3, column=0, padx=5, pady=5)

# Rótulo para o número de partes
label_parts = ttk.Label(frame, text="Número de partes:")
label_parts.grid(row=4, column=0, padx=5, pady=5)
entry_parts = ttk.Entry(frame)
entry_parts.grid(row=4, column=1, padx=5, pady=5)

# Botão para selecionar a pasta de destino
btn_select_output_directory = ttk.Button(frame, text="Selecionar Pasta de Destino", command=select_output_directory)
btn_select_output_directory.grid(row=5, column=0, padx=5, pady=5)

# Rótulo para exibir a pasta de destino selecionada
label_output_directory = ttk.Label(frame, text="")
label_output_directory.grid(row=6, column=0, padx=5, pady=5)

# Rótulo para exibir informações sobre as partes
label_parts_info = ttk.Label(frame, text="")
label_parts_info.grid(row=7, column=0, padx=5, pady=5)

# Botão para dividir o DOCX
btn_split = ttk.Button(frame, text="Dividir DOCX", command=lambda: split(input_docx, output_directory, int(entry_parts.get())))
btn_split.grid(row=8, column=0, columnspan=2, padx=5, pady=10)

root.mainloop()
